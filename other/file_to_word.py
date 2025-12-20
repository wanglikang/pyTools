#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将指定文件夹下符合特定后缀的文件内容写入到Word文档中
支持递归遍历文件夹
"""

import os
import argparse
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def file_to_word(folder_path, file_extensions, output_file):
    """
    将指定文件夹下符合特定后缀的文件内容写入到Word文档中
    
    Args:
        folder_path (str): 要遍历的文件夹路径
        file_extensions (list): 要匹配的文件后缀列表，如['.txt', '.py']
        output_file (str): 输出的Word文档路径
    """
    # 创建Word文档
    doc = Document()
    # 设置标题字号为小五
    title = doc.add_heading('文件内容汇总', 0)
    title_run = title.runs[0]
    title_run.font.size = Pt(9)
    
    # 设置页眉
    section = doc.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.text = "QLExpress辅助开发应用软件 V1.0"
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 设置页码 - "第 x 页 共 xx 页" 格式，居中对齐
    footer = section.footer
    # 清空现有页脚内容
    for paragraph in footer.paragraphs:
        paragraph.clear()
    
    # 创建新的页脚段落
    footer_paragraph = footer.add_paragraph()
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 构建页码格式
    # 在python-docx中，我们需要使用更复杂的方式添加字段
    # 先添加静态文本
    footer_paragraph.add_run("第 ")
    
    # 添加页码字段 (PAGE)
    # 注意：python-docx中add_field的正确用法
    from docx.oxml.shared import qn
    from docx.oxml import OxmlElement
    
    # 创建页码字段
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    # 将字段元素添加到run中
    r = footer_paragraph.add_run()
    r._r.append(fldChar1)
    r._r.append(instrText)
    r._r.append(fldChar2)
    
    # 添加中间文本
    footer_paragraph.add_run(" 页 共 ")
    
    # 添加总页数字段 (NUMPAGES)
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'begin')
    
    instrText2 = OxmlElement('w:instrText')
    instrText2.text = 'NUMPAGES'
    
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    
    # 将字段元素添加到run中
    r2 = footer_paragraph.add_run()
    r2._r.append(fldChar3)
    r2._r.append(instrText2)
    r2._r.append(fldChar4)
    
    # 添加结尾文本
    footer_paragraph.add_run(" 页")
    
    # 递归遍历文件夹
    file_count = 0
    line_count = 0
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            # 检查文件后缀
            if any(file.endswith(ext) for ext in file_extensions):
                file_path = os.path.join(root, file)
                file_count += 1
                
                try:
                    # 添加文件信息
                    relative_path = os.path.relpath(file_path, folder_path)
                    # doc.add_heading(f'文件 {file_count}: {relative_path}', level=1)
                    
                    # 读取文件内容并过滤空行
                    with open(file_path, 'r', encoding='utf-8') as f:
                        lines = f.readlines()
                    
                    # 过滤空行
                    non_empty_lines = [line.rstrip('\n') for line in lines if line.strip()]
                    line_count += len(non_empty_lines)
                    content = '\n'.join(non_empty_lines)
                    
                    # 添加文件内容，设置二级标题和正文内容的字号为小五
                    heading = doc.add_heading(f'文件:{file}', level=2)
                    # 设置二级标题字号
                    for run in heading.runs:
                        run.font.size = Pt(9)
                    
                    # 设置正文内容字号
                    paragraph = doc.add_paragraph()
                    paragraph_run = paragraph.add_run(content)
                    paragraph_run.font.size = Pt(9)
                    
                    # 添加分页符
                    doc.add_page_break()
                    
                    print(f'已处理文件: {relative_path}')
                except Exception as e:
                    doc.add_heading('内容读取失败:', level=2)
                    doc.add_paragraph(f'错误信息: {str(e)}')
                    doc.add_page_break()
                    print(f'处理文件失败: {relative_path}, 错误: {str(e)}')
    
    # 保存文档
    doc.save(output_file)
    print(f'\n处理完成！共处理 {file_count} 个文件')
    print(f'共处理 {line_count} 行内容')
    print(f'输出文件: {output_file}')

def main():
    """主函数"""
    parser = argparse.ArgumentParser(description='将指定文件夹下符合特定后缀的文件内容写入到Word文档中')
    parser.add_argument('-f', '--folder', required=True, help='要遍历的文件夹路径')
    parser.add_argument('-e', '--extensions', required=True, help='要匹配的文件后缀，多个后缀用逗号分隔，如：.txt,.py,.md')
    parser.add_argument('-o', '--output', default='file_content.docx', help='输出的Word文档路径，默认：file_content.docx')
    
    args = parser.parse_args()
    
    # 处理文件后缀参数
    extensions = [ext.strip() for ext in args.extensions.split(',')]
    
    # 验证文件夹路径
    if not os.path.exists(args.folder):
        print(f'错误：文件夹路径不存在：{args.folder}')
        return
    
    # 执行转换
    file_to_word(args.folder, extensions, args.output)

if __name__ == '__main__':
    main()